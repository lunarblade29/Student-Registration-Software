import time
from datetime import datetime
from flask import Flask, render_template, request, send_from_directory, jsonify
import pandas as pd
import os
from docx import Document  # Import python-docx
from werkzeug.utils import secure_filename
from my_dashboard.dashboard import dashboard_bp
from my_dashboard.dashboard import main_bp  # Adjust path as needed
import io

app = Flask(__name__)
app.register_blueprint(dashboard_bp)
app.register_blueprint(main_bp)  # This handles the homepage (/)

# Excel file path
EXCEL_FILE = "students.xlsx"
TEMPLATE_PATH = "templates/template.docx"  # Path to the Word template

REGISTRATION_DATA_FILE = "REGISTRATION_DATA_FILE.xlsx"
UPLOAD_FOLDER = "."  # Current directory for simplicity, adjust as needed
ALLOWED_EXTENSIONS = {"xlsx"}

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

today_date = datetime.today().strftime("%d %B %Y")

uploaded_file_data = None


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


# Ensure the Excel file exists
if not os.path.exists(EXCEL_FILE):
    df = pd.DataFrame(
        columns=[
            "Registration Number",
            "Name",
            "Email",
            "Biometric Verification Completed",
            "Gender",
            "Recent Photographs Submitted",
            "Marksheet X Verification",
            "Marksheet XII Verification",
            "Duration of Degree Course",
            "Degree Marksheet Verification",
            "Degree Certificate Verification",
            "Provisional Degree Certificate",
            "Graduation Marks Eligibility",
            "Graduation Status",
            "Incomplete Graduation Certificate",
            "Graduation Area",
            "Work Experience Certificate Submitted",
            "Work Experience Reason",
            "SC/ST/OBC/PwD Certificate Verification",
            "PwD Enclosure Submitted",
            "Demand Draft Submitted",
            "Demand Draft Amount",
            "Draft No.",
            "DT No.",
            "CAT Score Card Verification",
            "Guardian's Declaration Submitted",
            "Medical Info Form Submitted",
            "Personal Data Card Submitted",
            "Campus Rules Declaration Submitted",
            "Anti-Ragging Form Submitted",
            "Bank Details Declaration Submitted",
            "Bank Loan Taken",
            "Bank Name",
            "Loan Amount",
            "Remarks",
            "reason_photos",
            "reason_marksheet_x",
            "reason_marksheet_xii",
            "reason_degree_marksheet",
            "reason_degree_certificate",
            "reason_provisional_degree",
            "reason_incomplete_graduation",
            "reason_category_certificate",
            "reason_pwd_enclosure",
            "reason_cat_score_card",
            "reason_guardians_declaration",
            "reason_medical_info_form",
            "reason_personal_data_card",
            "reason_campus_rules_declaration",
            "reason_anti_ragging_form",
            "reason_bank_details_declaration",
        ]
    )
    df.to_excel(EXCEL_FILE, index=False)

# Ensure the registration data Excel file exists and has the 'Email' column
if not os.path.exists(REGISTRATION_DATA_FILE):
    reg_df = pd.DataFrame(columns=["Registration Number", "Name", "Email"])
    reg_df.to_excel(REGISTRATION_DATA_FILE, index=False)
else:
    reg_df = pd.read_excel(REGISTRATION_DATA_FILE)
    if "Email" not in reg_df.columns:
        reg_df["Email"] = ""  # Add Email column if it doesn't exist
        reg_df.to_excel(REGISTRATION_DATA_FILE, index=False)


@app.route("/upload_registration_data", methods=["POST"])
def upload_registration_data():
    global uploaded_file_data
    if "file" not in request.files:
        return "", 400  # Bad Request
    file = request.files["file"]
    if file.filename == "":
        return "", 400  # Bad Request
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        if filename.lower() == REGISTRATION_DATA_FILE.lower():
            try:
                uploaded_file_data = file.read()
                # Attempt a basic read to catch potential non-Excel files early
                pd.read_excel(io.BytesIO(uploaded_file_data))
                return "", 204  # No Content - Success (no message)
            except pd.errors.EmptyDataError:
                return (
                    "",
                    400,
                    {"ContentType": "text/plain"},
                )  # Bad Request - Empty Excel
            except Exception as e:
                print(f"Error during initial file read: {e}")
                return (
                    "",
                    400,
                    {"ContentType": "text/plain"},
                )  # Bad Request - Could not read Excel
        else:
            return "", 400  # Bad Request - Incorrect filename
    return "", 415  # Unsupported Media Type


@app.route("/finalize_upload_registration_data", methods=["POST"])
def finalize_upload_registration_data():
    global uploaded_file_data
    if uploaded_file_data:
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], REGISTRATION_DATA_FILE)
        try:
            with open(filepath, "wb") as f:
                f.write(uploaded_file_data)
            uploaded_file_data = None
            pd.read_excel(filepath)
            return "", 204  # No Content - Success
        except pd.errors.EmptyDataError:
            return (
                "",
                500,
                {"ContentType": "text/plain"},
            )  # Internal Server Error - Empty after save
        except Exception as e:
            print(f"Error saving or validating final file: {e}")
            return "", 500, {"ContentType": "text/plain"}  # Internal Server Error
    else:
        return (
            "",
            400,
            {"ContentType": "text/plain"},
        )  # Bad Request - Incorrect filename


@app.route("/get_student_info", methods=["POST"])
def get_student_info():
    reg_no = request.form.get("reg_no")
    if reg_no:
        try:
            reg_df = pd.read_excel(REGISTRATION_DATA_FILE)
            student_data = (
                reg_df[reg_df["Registration Number"] == reg_no].iloc[0].to_dict()
            )
            return jsonify(
                {
                    "name": student_data.get("Name", ""),
                    "email": student_data.get("Email", ""),
                }
            )
        except IndexError:
            return jsonify(
                {"name": "", "email": ""}
            )  # Return empty if registration number not found
        except Exception as e:
            print(f"Error reading registration data: {e}")
            return jsonify({"name": "", "email": ""})
    return jsonify({"name": "", "email": ""})


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        try:
            # Collect form data
            student_data = {
                "Name": request.form.get("name"),
                "Registration Number": request.form.get("reg_no"),
                "Email": request.form.get("email"),
                "Biometric Verification Completed": request.form.get("biometric"),
                "Gender": request.form.get("gender"),
                "Recent Photographs Submitted": request.form.get("photos"),
                "Marksheet X Verification": request.form.get("marksheet_x"),
                "Marksheet XII Verification": request.form.get("marksheet_xii"),
                "Duration of Degree Course": request.form.get("degree_duration"),
                "Degree Marksheet Verification": request.form.get("degree_marksheet"),
                "Degree Certificate Verification": request.form.get(
                    "degree_certificate"
                ),
                "Provisional Degree Certificate": request.form.get(
                    "provisional_degree"
                ),
                "Graduation Marks Eligibility": request.form.get("marks_obtained"),
                "Graduation Status": request.form.get("graduation_status"),
                "Incomplete Graduation Certificate": request.form.get(
                    "incomplete_graduation_cert"
                ),
                "Graduation Area": request.form.get("graduation_area"),
                "Work Experience Certificate Submitted": request.form.get(
                    "work_experience_certificate"
                ),
                "Work Experience Reason": request.form.get("reason_for_no"),
                "SC/ST/OBC/PwD Certificate Verification": request.form.get(
                    "category_certificate"
                ),
                "PwD Enclosure Submitted": request.form.get("pwd_enclosure"),
                "Demand Draft Submitted": request.form.get("demand_draft_submitted"),
                "Demand Draft Amount": (
                    request.form.get("demand_draft_amount")
                    if request.form.get("demand_draft_submitted") == "Yes"
                    else None
                ),
                "Draft No.": (
                    request.form.get("draft_no")
                    if request.form.get("demand_draft_submitted") == "Yes"
                    else None
                ),
                "DT No.": (
                    request.form.get("dt_no")
                    if request.form.get("demand_draft_submitted") == "Yes"
                    else None
                ),
                "CAT Score Card Verification": request.form.get("cat_score_card"),
                "Guardian's Declaration Submitted": request.form.get(
                    "guardians_declaration"
                ),
                "Medical Info Form Submitted": request.form.get("medical_info_form"),
                "Personal Data Card Submitted": request.form.get("personal_data_card"),
                "Campus Rules Declaration Submitted": request.form.get(
                    "campus_rules_declaration"
                ),
                "Anti-Ragging Form Submitted": request.form.get("anti_ragging_form"),
                "Bank Details Declaration Submitted": request.form.get(
                    "bank_details_declaration"
                ),
                "Bank Loan Taken": request.form.get("bank_loan"),
                "Bank Name": request.form.get("bank_name"),
                "Loan Amount": request.form.get("loan_amount"),
                "Remarks": request.form.get("remarks"),
                "reason_photos": request.form.get("reason_photos"),
                "reason_marksheet_x": request.form.get("reason_marksheet_x"),
                "reason_marksheet_xii": request.form.get("reason_marksheet_xii"),
                "reason_degree_marksheet": request.form.get("reason_degree_marksheet"),
                "reason_degree_certificate": request.form.get(
                    "reason_degree_certificate"
                ),
                "reason_provisional_degree": request.form.get(
                    "reason_provisional_degree"
                ),
                "reason_incomplete_graduation": request.form.get(
                    "reason_incomplete_graduation"
                ),
                "reason_category_certificate": request.form.get(
                    "reason_category_certificate"
                ),
                "reason_pwd_enclosure": request.form.get("reason_pwd_enclosure"),
                "reason_cat_score_card": request.form.get("reason_cat_score_card"),
                "reason_guardians_declaration": request.form.get(
                    "reason_guardians_declaration"
                ),
                "reason_medical_info_form": request.form.get(
                    "reason_medical_info_form"
                ),
                "reason_personal_data_card": request.form.get(
                    "reason_personal_data_card"
                ),
                "reason_campus_rules_declaration": request.form.get(
                    "reason_campus_rules_declaration"
                ),
                "reason_anti_ragging_form": request.form.get(
                    "reason_anti_ragging_form"
                ),
                "reason_bank_details_declaration": request.form.get(
                    "reason_bank_details_declaration"
                ),
            }

            # Append data to Excel
            df = pd.read_excel(EXCEL_FILE)
            df = df._append(student_data, ignore_index=True)
            df.to_excel(EXCEL_FILE, index=False)

            # 2. Update the Word Document
            # Load template and replace placeholders with form data
            doc = Document(
                "templates/template.docx"
            )  # Make sure the correct path is used

            for paragraph in doc.paragraphs:
                if "{name}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "{name}", student_data["Name"]
                    )
                if "{reg_no}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "{reg_no}", student_data["Registration Number"]
                    )
                if "{year}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "{year}", student_data["Duration of Degree Course"]
                    )
                if "{grad_status}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "{grad_status}", student_data["Graduation Status"]
                    )
                if "{resign}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "{resign}", student_data["Work Experience Reason"]
                    )
                if "{bank_name}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "{bank_name}", student_data["Bank Name"]
                    )
                if "{bank_amount}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "{bank_amount}", student_data["Loan Amount"]
                    )
                if "{date}" in paragraph.text:
                    paragraph.text = paragraph.text.replace("{date}", today_date)
                if "{remarks}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "{remarks}",
                        student_data["Remarks"]
                        + " "
                        + student_data["reason_photos"]
                        + " "
                        + student_data["reason_marksheet_x"]
                        + " "
                        + student_data["reason_marksheet_xii"]
                        + " "
                        + student_data["reason_degree_marksheet"]
                        + " "
                        + student_data["reason_degree_certificate"]
                        + " "
                        + student_data["reason_provisional_degree"]
                        + " "
                        + student_data["reason_incomplete_graduation"]
                        + " "
                        + student_data["reason_category_certificate"]
                        + " "
                        + student_data["reason_pwd_enclosure"]
                        + " "
                        + student_data["reason_cat_score_card"]
                        + " "
                        + student_data["reason_guardians_declaration"]
                        + " "
                        + student_data["reason_medical_info_form"]
                        + " "
                        + student_data["reason_personal_data_card"]
                        + " "
                        + student_data["reason_campus_rules_declaration"]
                        + " "
                        + student_data["reason_anti_ragging_form"]
                        + " "
                        + student_data["reason_bank_details_declaration"],
                    )

                # Handling the genders
                if "{gm}" in paragraph.text:
                    if student_data["Gender"] == "Male":
                        paragraph.text = paragraph.text.replace(
                            "{gm}", "☑"
                        )  # Tick for Yes
                        paragraph.text = paragraph.text.replace(
                            "{gf}", "☐"
                        )  # blank for No
                        paragraph.text = paragraph.text.replace(
                            "{go}", "☐"
                        )  # blank for No
                    elif student_data["Gender"] == "Female":
                        paragraph.text = paragraph.text.replace(
                            "{gm}", "☐"
                        )  # blank for Yes
                        paragraph.text = paragraph.text.replace(
                            "{gf}", "☑"
                        )  # Tick for N
                        paragraph.text = paragraph.text.replace(
                            "{go}", "☐"
                        )  # blank for No
                    else:
                        paragraph.text = paragraph.text.replace(
                            "{gm}", "☐"
                        )  # blank for Yes
                        paragraph.text = paragraph.text.replace(
                            "{gf}", "☐"
                        )  # Tick for N
                        paragraph.text = paragraph.text.replace(
                            "{go}", "☑"
                        )  # blank for No

                # Handling the drafts
                if "{dty}" in paragraph.text:
                    if student_data["Demand Draft Amount"] == "440000":
                        paragraph.text = paragraph.text.replace(
                            "{dty}", "☑"
                        )  # Tick for Yes
                        paragraph.text = paragraph.text.replace(
                            "{dtn}", "☐"
                        )  # blank for No
                        paragraph.text = paragraph.text.replace(
                            "{ddt}", student_data["Draft No."]
                        )  # Tick for Yes
                        paragraph.text = paragraph.text.replace(
                            "{dtt}", student_data["DT No."]
                        )  # blank for No
                    else:
                        paragraph.text = paragraph.text.replace(
                            "{dty}", "☐"
                        )  # blank for Yes
                        paragraph.text = paragraph.text.replace(
                            "{dtn}", "☑"
                        )  # Tick for N
                        paragraph.text = paragraph.text.replace(
                            "{ddt}", "------"
                        )  # Tick for Yes
                        paragraph.text = paragraph.text.replace(
                            "{dtt}", "------"
                        )  # blank for No
                if "{dcy}" in paragraph.text:
                    if student_data["Demand Draft Amount"] == "20000":
                        paragraph.text = paragraph.text.replace(
                            "{dcy}", "☑"
                        )  # Tick for Yes
                        paragraph.text = paragraph.text.replace(
                            "{dcn}", "☐"
                        )  # blank for No
                        paragraph.text = paragraph.text.replace(
                            "{ddc}", student_data["Draft No."]
                        )  # Tick for Yes
                        paragraph.text = paragraph.text.replace(
                            "{dtc}", student_data["DT No."]
                        )  # blank for No
                    else:
                        paragraph.text = paragraph.text.replace(
                            "{dcy}", "☐"
                        )  # blank for Yes
                        paragraph.text = paragraph.text.replace(
                            "{dcn}", "☑"
                        )  # Tick for N
                        paragraph.text = paragraph.text.replace(
                            "{ddc}", "------"
                        )  # Tick for Yes
                        paragraph.text = paragraph.text.replace(
                            "{dtc}", "------"
                        )  # blank for No
                if "{ddy}" in paragraph.text:
                    if student_data["Demand Draft Amount"] == "460000":
                        paragraph.text = paragraph.text.replace(
                            "{ddy}", "☑"
                        )  # Tick for Yes
                        paragraph.text = paragraph.text.replace(
                            "{ddn}", "☐"
                        )  # blank for No
                        paragraph.text = paragraph.text.replace(
                            "{ddb}", student_data["Draft No."]
                        )  # Tick for Yes
                        paragraph.text = paragraph.text.replace(
                            "{dtb}", student_data["DT No."]
                        )  # blank for No
                    else:
                        paragraph.text = paragraph.text.replace(
                            "{ddy}", "☐"
                        )  # blank for Yes
                        paragraph.text = paragraph.text.replace(
                            "{ddn}", "☑"
                        )  # Tick for N
                        paragraph.text = paragraph.text.replace(
                            "{ddb}", "------"
                        )  # Tick for Yes
                        paragraph.text = paragraph.text.replace(
                            "{dtb}", "------"
                        )  # blank for No

                # Handling the streams
                if (
                    "{gse}" in paragraph.text
                    and student_data["Graduation Area"] == "Engineering"
                ):
                    paragraph.text = paragraph.text.replace(
                        "{gse}", "☑"
                    )  # Tick for Yes
                    paragraph.text = paragraph.text.replace(
                        "{gss}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gsc}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gsa}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gso}", "☐"
                    )  # blank for No
                if (
                    "{gss}" in paragraph.text
                    and student_data["Graduation Area"] == "Science"
                ):
                    paragraph.text = paragraph.text.replace(
                        "{gss}", "☑"
                    )  # Tick for Yes
                    paragraph.text = paragraph.text.replace(
                        "{gse}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gsc}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gsa}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gso}", "☐"
                    )  # blank for No
                if (
                    "{gsc}" in paragraph.text
                    and student_data["Graduation Area"] == "Commerce"
                ):
                    paragraph.text = paragraph.text.replace(
                        "{gsc}", "☑"
                    )  # Tick for Yes
                    paragraph.text = paragraph.text.replace(
                        "{gss}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gse}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gsa}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gso}", "☐"
                    )  # blank for No
                if (
                    "{gsa}" in paragraph.text
                    and student_data["Graduation Area"] == "Arts"
                ):
                    paragraph.text = paragraph.text.replace(
                        "{gsa}", "☑"
                    )  # Tick for Yes
                    paragraph.text = paragraph.text.replace(
                        "{gss}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gsc}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gse}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gso}", "☐"
                    )  # blank for No
                if (
                    "{gso}" in paragraph.text
                    and student_data["Graduation Area"] == "Others"
                ):
                    paragraph.text = paragraph.text.replace(
                        "{gso}", "☑"
                    )  # Tick for Yes
                    paragraph.text = paragraph.text.replace(
                        "{gss}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gsc}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gsa}", "☐"
                    )  # blank for No
                    paragraph.text = paragraph.text.replace(
                        "{gse}", "☐"
                    )  # blank for No

            # Helper function to handle Yes/No placeholder replacement
            def handle_checkbox_placeholder(
                paragraph,
                placeholder_yes,
                placeholder_no,
                placeholder_na,
                student_data_value,
            ):
                if (
                    placeholder_yes in paragraph.text
                    or placeholder_no in paragraph.text
                    or placeholder_na in paragraph.text
                ):
                    if student_data_value == "Yes":
                        paragraph.text = paragraph.text.replace(
                            placeholder_yes, "☑"
                        )  # Tick for Yes
                        paragraph.text = paragraph.text.replace(
                            placeholder_no, "☐"
                        )  # blank for No
                        paragraph.text = paragraph.text.replace(
                            placeholder_na, "☐"
                        )  # blank for NA
                    elif student_data_value == "No":
                        paragraph.text = paragraph.text.replace(
                            placeholder_yes, "☐"
                        )  # blank for Yes
                        paragraph.text = paragraph.text.replace(
                            placeholder_no, "☑"
                        )  # Tick for No
                        paragraph.text = paragraph.text.replace(
                            placeholder_na, "☐"
                        )  # blank for NA
                    else:
                        paragraph.text = paragraph.text.replace(
                            placeholder_yes, "☐"
                        )  # blank for Yes
                        paragraph.text = paragraph.text.replace(
                            placeholder_no, "☐"
                        )  # blank for No
                        paragraph.text = paragraph.text.replace(
                            placeholder_na, "☑"
                        )  # Tick for NA

            # Now use this helper function for all fields
            for paragraph in doc.paragraphs:
                handle_checkbox_placeholder(
                    paragraph,
                    "{by}",
                    "{bn}",
                    "{na}",
                    student_data["Biometric Verification Completed"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{py}",
                    "{pn}",
                    "{na}",
                    student_data["Recent Photographs Submitted"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{xy}",
                    "{xn}",
                    "{na}",
                    student_data["Marksheet X Verification"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{xiiy}",
                    "{xiin}",
                    "{na}",
                    student_data["Marksheet XII Verification"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{my}",
                    "{mn}",
                    "{na}",
                    student_data["Degree Marksheet Verification"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{cy}",
                    "{cn}",
                    "{na}",
                    student_data["Degree Certificate Verification"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{gy}",
                    "{gn}",
                    "{na}",
                    student_data["Graduation Marks Eligibility"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{csy}",
                    "{csn}",
                    "{na}",
                    student_data["CAT Score Card Verification"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{gdy}",
                    "{gdn}",
                    "{na}",
                    student_data["Guardian's Declaration Submitted"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{miy}",
                    "{min}",
                    "{na}",
                    student_data["Medical Info Form Submitted"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{pcy}",
                    "{pcn}",
                    "{na}",
                    student_data["Personal Data Card Submitted"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{cry}",
                    "{crn}",
                    "{na}",
                    student_data["Campus Rules Declaration Submitted"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{ary}",
                    "{arn}",
                    "{na}",
                    student_data["Anti-Ragging Form Submitted"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{bdy}",
                    "{bdn}",
                    "{na}",
                    student_data["Bank Details Declaration Submitted"],
                )
                handle_checkbox_placeholder(
                    paragraph, "{loy}", "{lon}", "{na}", student_data["Bank Loan Taken"]
                )
                # 3 fields
                handle_checkbox_placeholder(
                    paragraph,
                    "{pdy}",
                    "{pdn}",
                    "{pdna}",
                    student_data["Provisional Degree Certificate"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{igy}",
                    "{ign}",
                    "{igna}",
                    student_data["Incomplete Graduation Certificate"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{ry}",
                    "{rn}",
                    "{rna}",
                    student_data["Work Experience Certificate Submitted"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{ccy}",
                    "{ccn}",
                    "{ccna}",
                    student_data["SC/ST/OBC/PwD Certificate Verification"],
                )
                handle_checkbox_placeholder(
                    paragraph,
                    "{psy}",
                    "{psn}",
                    "{psna}",
                    student_data["PwD Enclosure Submitted"],
                )

            # Save the updated document
            time.sleep(0.1)
            output_path = os.path.join("downloads", "output.docx")
            doc.save(output_path)

            # Example success response
            return jsonify({"message": "Form submitted successfully!"})

        except Exception as e:
            print("Error:", str(e))  # Print the error in the console
            return jsonify({"error": "Internal Server Error"}), 500

    return render_template("index.html")


@app.route("/download")
def download():
    return send_from_directory("downloads", "output.docx", as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True)
